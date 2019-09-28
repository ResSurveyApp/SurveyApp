
from flask import Flask, request, render_template,jsonify
import requests
import businessLogic as bl
import datetime
import os
import pandas as pd
from pandas import ExcelWriter
from pandas import ExcelFile
from shutil import copyfile
import chartsLogic as cl
from flask_mail import Mail, Message
import warnings
import time
import json
import numpy as np
import functools as ft
from bson import json_util
import xlrd
import pprint
import matplotlib.pyplot as plt
import xlwt
from xlwt import Workbook

import base64
warnings.filterwarnings("ignore")
import smtplib
from flask_cors import CORS
from bson.json_util import dumps,loads
from PyPDF2 import PdfFileReader, PdfFileWriter
from docxtpl import DocxTemplate
from docx.shared import Inches


warnings.filterwarnings("ignore")

sectors=['R1','R2','R3','R4']
subsectors=['Physical','Organizational','Technical']
df=pd.DataFrame()
app = Flask(__name__)

APP_ROOT = os.path.dirname(os.path.abspath(__file__))
cors = CORS(app)

@app.route('/backend')
def login():


    return "<u><b>Login Page</u></b>"
## Redirect to login page

@app.route('/backend/questionsUpload', methods=['GET','POST'])
def questionsUpload():
    survey = request.args.get('survey')
    company = request.args.get('company')
    file = request.files['file']
    #file = xlrd.open_workbook("C:/Users/headway/PycharmProjects/untitled/questions.xlsx")
    df = pd.read_excel(file)
    n=1


    df1,n=bl.parse(df,'R1','Physical',n)
    df2,n=bl.parse(df,'R1','Organizational',n)
    df3,n=bl.parse(df,'R1','Technical',n)
    df4,n=bl.parse(df,'R2','Physical',n)
    df5,n=bl.parse(df,'R2','Organizational',n)
    df6,n=bl.parse(df,'R2','Technical',n)
    df7,n=bl.parse(df,'R3','Physical',n)
    df8,n=bl.parse(df,'R3','Organizational',n)
    df9,n=bl.parse(df,'R3','Technical',n)
    df10,n=bl.parse(df,'R4','Physical',n)
    df11,n=bl.parse(df,'R4','Organizational',n)
    df12,n=bl.parse(df,'R4','Technical',n)
    df=pd.concat([df1,df2,df3,df4,df5,df6,df7,df8,df9,df10,df11,df12])
    print(df1[['sector','subsector','category_id','category_name','qid']])
    print(df2[['sector','subsector','category_id','category_name','qid']])
    print(df3[['sector','subsector','category_id','category_name','qid']])
    print(df4[['sector','subsector','category_id','category_name','qid']])
    print(df5[['sector','subsector','category_id','category_name','qid']])
    print(df6[['sector','subsector','category_id','category_name','qid']])
    print(df7[['sector','subsector','category_id','category_name','qid']])
    print(df8[['sector','subsector','category_id','category_name','qid']])
    print(df9[['sector','subsector','category_id','category_name','qid']])
    print(df10[['sector','subsector','category_id','category_name','qid']])
    print(df11[['sector','subsector','category_id','category_name','qid']])
    print(df12[['sector','subsector','category_id','category_name','qid']])
    print("MAINFRAME")
    print(df[['sector','subsector','category_name','category_id','qid']])
    fileName='dataframe_BK.xlsx'
    wb = Workbook()
    sheet1 = wb.add_sheet('Sheet 1')
    wb.save(fileName)

    writer = ExcelWriter(fileName)
    df.to_excel(writer,'Sheet1',index=False)
    writer.save()
    
    #df.pop('category_id')

    #print(df)
    #Getting timestamp
    backup=bl.getTimeStamp()

    #code pending for taking backup of this json
    jFile= './db.json'
    jsonFile= './db_{}.json'.format(backup)
    copyfile(jFile, jsonFile)

    #Parsing and divding into sector chunks of data from excel
    df_R1 = df[df['sector']=='R1']
    df_R2 = df[df['sector']=='R2']
    df_R3 = df[df['sector']=='R3']
    df_R4 = df[df['sector']=='R4']

    #initialize Mongo DB details or can be taken from a config file in future
    host='localhost'
    database='Surveyapp'
    collection='questions'
    user='root'
    pwd='Blackpearl'


    #update the db.json file with related info
    bl.parseDF(df_R1,df_R2,df_R3,df_R4,jsonFile,survey,company)

    #push json to mongo
    bl.pushMongoDB(host,database,collection,jsonFile,user,pwd,survey,company)

    #function to clean the workspace and close all handlers and files
    #bl.cleanup()

    return 'Upload Done -- replace this with web page'

@app.route('/backend/usersUpload', methods=['GET','POST'])
def usersUpload():
    survey = request.args.get('survey')
    company = request.args.get('company')
    file = request.files['file']
    #file = xlrd.open_workbook("/apps/surveyapp_python/user_details.xlsx")
    df = pd.read_excel(file)

    # Getting timestamp
    backup = bl.getTimeStamp()

    # code pending for taking backup of this json
    jFile = './users.json'
    jsonFile = './users_{}.json'.format(backup)
    copyfile(jFile, jsonFile)
    lFile = './login.json'
    ljsonFile = './login_{}.json'.format(backup)
    copyfile(lFile, ljsonFile)
    # Parsing and divding into sector chunks of data from excel
    #df_username = df[df['username']]


    # initialize Mongo DB details or can be taken from a config file in future
    host = 'localhost'
    database = 'Surveyapp'
    collection = 'userdetails'
    user = 'root'
    pwd = 'Blackpearl'

    # update the db.json file with related info
    bl.parseDFusers(df, jsonFile, ljsonFile, survey, company)

    # push json to mongo
    bl.pushMongoDB(host, database, collection, jsonFile, user, pwd, survey, company)

    # function to clean the workspace and close all handlers and files
    # bl.cleanup()

    return 'Upload Done -- replace this with web page'

@app.route('/backend/companyNames', methods=['GET', 'POST'])
def companyNames():
    companies_list = list(bl.getcompanynames())
    print(companies_list)
    return json.dumps(companies_list)

@app.route('/backend/companySurveyNames', methods=['GET', 'POST'])
def companySurveyNames():
    company = request.args.get('company')
    #company = "company3"
    companysurveynames=list(bl.getcompanysurveynames(company))
    return json.dumps(companysurveynames)
	
@app.route('/backend/questionCount', methods=['GET', 'POST'])
def questionCount():
    company = request.args.get('company')
    survey = request.args.get('survey')
    #company = "company3"
    count =bl.getquestioncount(company, survey, sectors, subsectors)
    print(count)
    return json.dumps(count)

@app.route('/backend/autoSavedResponse', methods=['GET', 'POST'])
def autoSavedResponse():
    company = request.args.get('company')
    survey = request.args.get('survey')
    email = request.args.get('email')
    #company = "company3"
    count =bl.getautoSavedResponse(company, survey, email)
    print(count)
    return json.dumps(count)

@app.route('/backend/departmentNames', methods=['GET', 'POST'])
def departmentNames():
    company = request.args.get('company')
    survey = request.args.get('survey')
    #company = "company3"
    #survey = "survey3"
    departmentnames=list(bl.getdepartmentnames(company, survey))
    return json.dumps(departmentnames)

@app.route('/backend/releaseSurvey', methods=['GET', 'POST'])
def releaseSurvey():
    survey = request.args.get('survey')
    company = request.args.get('company')
    department = request.args.get('department')
    host = request.args.get('host')
    #url = request.args.get('url')

    #survey = "survey"
    #company = "company"
    #department = 'Deco'
    # bl.activateSurvey(survey,company,department)
    #mail_list=["vigneshsubramani28@gmail.com"]
    mail_list=list(bl.getMails(survey,company,department))
    print(mail_list)
    sendmail(survey,company,mail_list,host)
    return 'sent email'

@app.route('/backend/surveyQuestions', methods=['GET', 'POST'])
def surveyQuestions():
    survey = request.args.get('survey')
    company = request.args.get('company')
    sector = request.args.get('sector')
    subsector = request.args.get('subsector')
    cname = request.args.get('cname')
    email = request.args.get('email')
    questions = bl.displayquestions(survey,company,sector,subsector,cname, email)
    return json.dumps(questions, default=json_util.default)


@app.route('/backend/categories', methods=['GET', 'POST'])
def categories():
    survey = request.args.get('survey')
    company = request.args.get('company')
    sector = request.args.get('sector')
    subsector = request.args.get('subsector')
    #survey = "survey"
    #company = "company"
    #sector = "R1"
    #subsector = "Physical"
    count = bl.getcategory(survey,company,sector,subsector)
    print(count)
    return json.dumps(count, default=json_util.default)



@app.route('/backend/getAllPie', methods=['GET'])
def getAllPie():

    dataframe = request.args.get('dataframe')
    df=pd.read_excel(dataframe)
    data=cl.pie(df)
    return data


@app.route('/backend/getAllBar', methods=['GET'])
def getAllBar():

    dataframe = request.args.get('dataframe')
    df=pd.read_excel(dataframe)
    data=cl.bar(df)
    return data

@app.route('/backend/validation', methods=['GET', 'POST'])
def validation():
    survey = request.args.get('survey')
    company = request.args.get('company')
    email = request.args.get('email')
    host = request.args.get('host')
    #survey = "survey"
    #company = "company"
    #email = "pranushajanapaa@gmail.com"
    return bl.validate(survey,company,email, host)

@app.route('/backend/Registration', methods=['GET', 'POST'])
def Registration():
    survey = request.args.get('survey')
    company = request.args.get('company')
    email = request.args.get('email')
    username =  request.args.get('username')
    password =  request.args.get('password')
    return bl.register(survey,company,email,username,password)

@app.route('/backend/getAllRadar', methods=['GET'])
def getAllRadar():
    dataframe = request.args.get('dataframe')
    df=pd.read_excel(dataframe)
    data=cl.radar(df)
    return data

@app.route('/backend/userResponse', methods=['GET', 'POST'])
def userResponse():
    data = request.get_json()
    return bl.userResponseload(data, sectors, subsectors)

@app.route('/backend/saveResponse', methods=['GET', 'POST'])
def saveResponse():
    data = request.get_json()
    return bl.saveUserResponse(data)


@app.route('/backend/getAllRadarBySector', methods=['GET'])
def getAllRadarBySector():

    dataframe = request.args.get('dataframe')
    df=pd.read_excel(dataframe)
    data=cl.radarBySector(df,sectors)
    return data


@app.route('/backend/getSingleRadarAllSectors', methods=['GET'])
def getSingleRadarAllSectors():

    survey = request.args.get('survey')
    company = request.args.get('company')
    userid= request.args.get('userid')
    r = requests.get("http://localhost/backend/chartsOne?survey={}&company={}&userid={}".format(survey,company,userid))
    jsonres=r.json()
    df=pd.DataFrame(jsonres)
    df=df.sort_values(['sector','cid','subsector'])

    data=cl.radarAllSectors(df,sectors)
    return data



@app.route('/backend/upload', methods=['GET', 'POST'])
def upload():
    folder_name = request.form['superhero']
    '''
    # this is to verify that folder to upload to exists.
    if os.path.isdir(os.path.join(APP_ROOT, 'files/{}'.format(folder_name))):
        print("folder exist")
    '''
    target = os.path.join(APP_ROOT, 'files/{}'.format(folder_name))
    print(target)
    if not os.path.isdir(target):
        os.mkdir(target)
    print(request.files.getlist("file"))
    for upload in request.files.getlist("file"):
        print(upload)
        print("{} is the file name".format(upload.filename))
        filename = upload.filename
        # This is to verify files are supported

        destination = "/".join([target, filename])
        print("Accept incoming file:", filename)
        print("Save it to:", destination)
        upload.save(destination)

    # return send_from_directory("images", filename, as_attachment=True)
    return render_template("complete.html", image_name=filename)




############################################################################################################################################
############################################################################################################################################

@app.route('/backend/chartsAll', methods=['GET'])
def chartsAll():
    import time
    import xlwt
    from xlwt import Workbook
    survey = request.args.get('survey')
    company = request.args.get('company')
    host,base,colection,dbuser,pwd=bl.mongoInit('users')
    # print(host, base, colection, dbuser, pwd, "Checkinggg", survey, company, sectors)
    df=pd.DataFrame(columns=['sector','subsector','cid','qid','qscore','qconfidence'])
    for sector in sectors:
        document= bl.getSurveyDetailsAll(survey,company,host,base,colection,dbuser,pwd,sector)

        for i in document:
            df=df.append({'sector': i['rows']['sector'],'subsector':i['rows']['subsector'],'cid':i['rows']['cid'],'cname':i['rows']['cname'],'qid':i['rows']['qid'],'qscore':i['rows']['qscore'], 'qconfidence':i['rows']['qconfidence']},ignore_index=True)
    df=df.sort_values(['sector','subsector'])


    df['qconfidence'] = df['qconfidence'].astype(int)
    df['qscore'] = df['qscore'].astype(int)

    R1_df=df[df['sector']=='R1']
    R2_df=df[df['sector']=='R2']
    R3_df=df[df['sector']=='R3']
    R4_df=df[df['sector']=='R4']

    R1_df=cl.calculate(R1_df,colection)
    R2_df=cl.calculate(R2_df,colection)
    R3_df=cl.calculate(R3_df,colection)
    R4_df=cl.calculate(R4_df,colection)

    framelist=[R1_df,R2_df,R3_df,R4_df]
    mf=pd.DataFrame(columns=R1_df.columns.values)
    for i in framelist:

        i['avgsectscore']=i['qscore'].sum()/len(i['qscore'])
        mf=mf.append(i)
        mf.index = pd.RangeIndex(len(mf.index))


        #phy=i[i['subsector']=='Physical']
        #org=i[i['subsector']=='Organizational']
        #tech=i[i['subsector']=='Technical']

        #i['avgsectscore']=(phy.iloc[0]['subsector_avg']+org.iloc[0]['subsector_avg']+tech.iloc[0]['subsector_avg'])/3
        #mf=mf.append(i)
        #mf.index = pd.RangeIndex(len(mf.index))

        #avgR1_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')])

    avgR1_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')])
    avgR1_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R1')])
    avgR1_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R1')])

    avgR2_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R2')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R2')])
    avgR2_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R2')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R2')])
    avgR2_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R2')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R2')])

    avgR3_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R3')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R3')])
    avgR3_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R3')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R3')])
    avgR3_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R3')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R3')])

    avgR4_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R4')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R4')])
    avgR4_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R4')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R4')])
    avgR4_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R4')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R4')])



   # mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')]['subsector_avg']=avgR1_phy
   ## mf['subsector_avg'] = np.where(  (mf['subsector']=='Physical') & (mf['sector']=='R1'))

    filter1 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')]
    filter1['subsector_avg']=avgR1_phy

    filter2 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R1')]
    filter2['subsector_avg']=avgR1_org

    filter3 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R1')]
    filter3['subsector_avg']=avgR1_tech

    filter4 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R2')]
    filter4['subsector_avg']=avgR2_phy

    filter5 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R2')]
    filter5['subsector_avg']=avgR2_org

    filter6 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R2')]
    filter6['subsector_avg']=avgR2_tech

    filter7 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R3')]
    filter7['subsector_avg']=avgR3_phy

    filter8 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R3')]
    filter8['subsector_avg']=avgR3_org

    filter9 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R3')]
    filter9['subsector_avg']=avgR3_tech

    filter10 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R4')]
    filter10['subsector_avg']=avgR4_phy

    filter11 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R4')]
    filter11['subsector_avg']=avgR4_org

    filter12 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R4')]
    filter12['subsector_avg']=avgR4_tech

    dfs=[filter1,filter2,filter3,filter4,filter5,filter6,filter7,filter8,filter9,filter10,filter11,filter12]
    df=pd.concat(dfs)

    time=int(time.time())
    uniqId=545 ## Sent from UI later.. can be admin userid
    fileName='dataframe_{}.xlsx'.format(time)
    wb = Workbook()
    sheet1 = wb.add_sheet('Sheet 1')
    wb.save(fileName)

    writer = ExcelWriter(fileName)
    df.to_excel(writer,'Sheet1',index=False)
    writer.save()

    return fileName



############################################################################################################################################
############################################################################################################################################







@app.route('/backend/chartsOne', methods=['GET'])
def chartsOne():

    survey = request.args.get('survey')
    company = request.args.get('company')
    userid= request.args.get('userid')
    host,base,colection,dbuser,pwd=bl.mongoInit('users')


    df=pd.DataFrame(columns=['sector','subsector','cid','qid','qscore','qconfidence'])
    for sector in sectors:
        document= bl.getSurveyDetails(userid,survey,company,host,base,colection,dbuser,pwd,sector)
        for i in document:
            df=df.append({'sector': i['rows']['sector'],'subsector':i['rows']['subsector'],'cid':i['rows']['cid'],'cname':i['rows']['cname'],'qid':i['rows']['qid'],'qscore':i['rows']['qscore'], 'qconfidence':i['rows']['qconfidence']},ignore_index=True)
    df=df.sort_values(['sector','subsector'])


    df['qconfidence'] = df['qconfidence'].astype(int)
    df['qscore'] = df['qscore'].astype(int)


    R1_df=df[df['sector']=='R1']
    R2_df=df[df['sector']=='R2']
    R3_df=df[df['sector']=='R3']
    R4_df=df[df['sector']=='R4']

    R1_df=cl.calculate(R1_df,colection)
    R2_df=cl.calculate(R2_df,colection)
    R3_df=cl.calculate(R3_df,colection)
    R4_df=cl.calculate(R4_df,colection)

    framelist=[R1_df,R2_df,R3_df,R4_df]
    mf=pd.DataFrame(columns=R1_df.columns.values)
    for i in framelist:

        i['avgsectscore']=i['qscore'].sum()/len(i['qscore'])
        mf=mf.append(i)
        mf.index = pd.RangeIndex(len(mf.index))


        #phy=i[i['subsector']=='Physical']
        #org=i[i['subsector']=='Organizational']
        #tech=i[i['subsector']=='Technical']

        #i['avgsectscore']=(phy.iloc[0]['subsector_avg']+org.iloc[0]['subsector_avg']+tech.iloc[0]['subsector_avg'])/3
        #mf=mf.append(i)
        #mf.index = pd.RangeIndex(len(mf.index))

        #avgR1_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')])

    avgR1_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')])
    avgR1_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R1')])
    avgR1_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R1')])

    avgR2_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R2')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R2')])
    avgR2_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R2')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R2')])
    avgR2_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R2')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R2')])

    avgR3_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R3')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R3')])
    avgR3_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R3')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R3')])
    avgR3_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R3')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R3')])

    avgR4_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R4')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R4')])
    avgR4_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R4')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R4')])
    avgR4_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R4')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R4')])



   # mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')]['subsector_avg']=avgR1_phy
   ## mf['subsector_avg'] = np.where(  (mf['subsector']=='Physical') & (mf['sector']=='R1'))

    filter1 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')]
    filter1['subsector_avg']=avgR1_phy

    filter2 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R1')]
    filter2['subsector_avg']=avgR1_org

    filter3 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R1')]
    filter3['subsector_avg']=avgR1_tech

    filter4 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R2')]
    filter4['subsector_avg']=avgR2_phy

    filter5 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R2')]
    filter5['subsector_avg']=avgR2_org

    filter6 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R2')]
    filter6['subsector_avg']=avgR2_tech

    filter7 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R3')]
    filter7['subsector_avg']=avgR3_phy

    filter8 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R3')]
    filter8['subsector_avg']=avgR3_org

    filter9 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R3')]
    filter9['subsector_avg']=avgR3_tech

    filter10 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R4')]
    filter10['subsector_avg']=avgR4_phy

    filter11 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R4')]
    filter11['subsector_avg']=avgR4_org

    filter12 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R4')]
    filter12['subsector_avg']=avgR4_tech

    dfs=[filter1,filter2,filter3,filter4,filter5,filter6,filter7,filter8,filter9,filter10,filter11,filter12]
    df=pd.concat(dfs)

    js=df.to_json()


    return js






############################################################################################################################################
############################################################################################################################################


@app.route('/backend/chartsByDept', methods=['GET'])
def chartsByDept():
    import time
    import xlwt
    from xlwt import Workbook
    survey = request.args.get('survey')
    company = request.args.get('company')
    dept = request.args.get('department')
    host,base,colection,dbuser,pwd=bl.mongoInit('users')

    df=pd.DataFrame(columns=['sector','subsector','cid','qid','qscore','qconfidence'])
    for sector in sectors:
        document= bl.getSurveyDetailsByDept(survey,company,host,base,colection,dbuser,pwd,sector,dept)

        for i in document:
            df=df.append({'sector': i['rows']['sector'],'subsector':i['rows']['subsector'],'cid':i['rows']['cid'],'cname':i['rows']['cname'],'qid':i['rows']['qid'],'qscore':i['rows']['qscore'], 'qconfidence':i['rows']['qconfidence']},ignore_index=True)
    df=df.sort_values(['sector','subsector'])


    df['qconfidence'] = df['qconfidence'].astype(int)
    df['qscore'] = df['qscore'].astype(int)

    R1_df=df[df['sector']=='R1']
    R2_df=df[df['sector']=='R2']
    R3_df=df[df['sector']=='R3']
    R4_df=df[df['sector']=='R4']

    R1_df=cl.calculate(R1_df,colection)
    R2_df=cl.calculate(R2_df,colection)
    R3_df=cl.calculate(R3_df,colection)
    R4_df=cl.calculate(R4_df,colection)

    framelist=[R1_df,R2_df,R3_df,R4_df]
    mf=pd.DataFrame(columns=R1_df.columns.values)
    for i in framelist:

        i['avgsectscore']=i['qscore'].sum()/len(i['qscore'])
        mf=mf.append(i)
        mf.index = pd.RangeIndex(len(mf.index))


        #phy=i[i['subsector']=='Physical']
        #org=i[i['subsector']=='Organizational']
        #tech=i[i['subsector']=='Technical']

        #i['avgsectscore']=(phy.iloc[0]['subsector_avg']+org.iloc[0]['subsector_avg']+tech.iloc[0]['subsector_avg'])/3
        #mf=mf.append(i)
        #mf.index = pd.RangeIndex(len(mf.index))

        #avgR1_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')])

    avgR1_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')])
    avgR1_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R1')])
    avgR1_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R1')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R1')])

    avgR2_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R2')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R2')])
    avgR2_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R2')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R2')])
    avgR2_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R2')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R2')])

    avgR3_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R3')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R3')])
    avgR3_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R3')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R3')])
    avgR3_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R3')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R3')])

    avgR4_phy=mf[(mf['subsector']=='Physical') & (mf['sector']=='R4')].qscore.sum()/len(mf[(mf['subsector']=='Physical') & (mf['sector']=='R4')])
    avgR4_org=mf[(mf['subsector']=='Organizational') & (mf['sector']=='R4')].qscore.sum()/len(mf[(mf['subsector']=='Organizational') & (mf['sector']=='R4')])
    avgR4_tech=mf[(mf['subsector']=='Technical') & (mf['sector']=='R4')].qscore.sum()/len(mf[(mf['subsector']=='Technical') & (mf['sector']=='R4')])



   # mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')]['subsector_avg']=avgR1_phy
   ## mf['subsector_avg'] = np.where(  (mf['subsector']=='Physical') & (mf['sector']=='R1'))

    filter1 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R1')]
    filter1['subsector_avg']=avgR1_phy

    filter2 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R1')]
    filter2['subsector_avg']=avgR1_org

    filter3 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R1')]
    filter3['subsector_avg']=avgR1_tech

    filter4 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R2')]
    filter4['subsector_avg']=avgR2_phy

    filter5 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R2')]
    filter5['subsector_avg']=avgR2_org

    filter6 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R2')]
    filter6['subsector_avg']=avgR2_tech

    filter7 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R3')]
    filter7['subsector_avg']=avgR3_phy

    filter8 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R3')]
    filter8['subsector_avg']=avgR3_org

    filter9 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R3')]
    filter9['subsector_avg']=avgR3_tech

    filter10 = mf[(mf['subsector']=='Physical') & (mf['sector']=='R4')]
    filter10['subsector_avg']=avgR4_phy

    filter11 = mf[(mf['subsector']=='Organizational') & (mf['sector']=='R4')]
    filter11['subsector_avg']=avgR4_org

    filter12 = mf[(mf['subsector']=='Technical') & (mf['sector']=='R4')]
    filter12['subsector_avg']=avgR4_tech

    dfs=[filter1,filter2,filter3,filter4,filter5,filter6,filter7,filter8,filter9,filter10,filter11,filter12]
    df=pd.concat(dfs)

    time=int(time.time())
    uniqId=545 ## Sent from UI later.. can be admin userid
    fileName='dataframe_{}.xlsx'.format(time)
    wb = Workbook()
    sheet1 = wb.add_sheet('Sheet 1')
    wb.save(fileName)

    writer = ExcelWriter(fileName)
    df.to_excel(writer,'Sheet1',index=False)
    writer.save()

    return fileName

############################################################################################################################################





############################################################################################################################################
############################################################################################################################################


@app.route('/backend/getSinglePie', methods=['GET'])
def getSinglePie():

    survey = request.args.get('survey')
    company = request.args.get('company')
    userid= request.args.get('userid')
    print(survey,company,userid)
    #r = requests.get("http://localhost/backend/chartsOne?survey='Quarter%201'&company='ITH'&userid='588'")
    r = requests.get("http://localhost/backend/chartsOne?survey={}&company={}&userid={}".format(survey,company,userid))


    jsonres=r.json()
    df=pd.DataFrame(jsonres)

    data=cl.pie(df)
    return data

@app.route('/backend/getTable', methods=['GET'])
def getTable():
    survey = request.args.get('survey')
    company = request.args.get('company')
    userid= request.args.get('userid')
    r = requests.get("http://localhost/backend/chartsOne?survey={}&company={}&userid={}".format(survey,company,userid))
    jsonres=r.json()
    df=pd.DataFrame(jsonres)
    data=cl.getTable(df,sectors,subsectors)
    return data

@app.route('/backend/getSingleBar', methods=['GET'])
def getSingleBar():

    survey = request.args.get('survey')
    company = request.args.get('company')
    userid= request.args.get('userid')

    r = requests.get("http://localhost/backend/chartsOne?survey={}&company={}&userid={}".format(survey,company,userid))
    jsonres=r.json()
    df=pd.DataFrame(jsonres)

    data=cl.bar(df)
    return data

@app.route('/backend/getAllRadarAllSectors', methods=['GET'])
def getAllRadarAllSectors():

    dataframe = request.args.get('dataframe')
    df=pd.read_excel(dataframe)
    data=cl.radarAllSectors(df,sectors)
    return data


@app.route('/backend/getSingleRadar', methods=['GET'])
def getSingleRadar():

    survey = request.args.get('survey')
    company = request.args.get('company')
    userid= request.args.get('userid')

    r = requests.get("http://localhost/backend/chartsOne?survey={}&company={}&userid={}".format(survey,company,userid))
    jsonres=r.json()
    df=pd.DataFrame(jsonres)

    data=cl.radar(df)
    return data




@app.route('/backend/getSingleRadarBySector', methods=['GET'])
def getSingleRadarBySector():

    survey = request.args.get('survey')
    company = request.args.get('company')
    userid= request.args.get('userid')
#    sector= request.args.get('sector')

    r = requests.get("http://localhost/backend/chartsOne?survey={}&company={}&userid={}".format(survey,company,userid))
    jsonres=r.json()
    df=pd.DataFrame(jsonres)

    data=cl.radarBySector(df,sectors)
    return data


@app.route('/backend/getSingleTable', methods=['GET'])
def getSingleTable():

    survey = request.args.get('survey')
    company = request.args.get('company')
    userid= request.args.get('userid')
    r = requests.get("http://localhost/backend/chartsOne?survey={}&company={}&userid={}".format(survey,company,userid))
    jsonres=r.json()
    df=pd.DataFrame(jsonres)
    df=df.sort_values(['sector','cid','subsector'])
    df=df[df['sector']=='R1']
    data=dumps(df)
    return data


@app.route('/backend/userlogin', methods=['GET', 'POST'])
def userlogin():
    data = request.get_json()
    jsonFile = './sessionTemplate.json'
    check = bl.userLogin(data, jsonFile)
    print(check)
    return json.dumps(check)


def sendmail(survey,company,mails, host):
    to = mails

    gmail_user = 'perugubharathkumar@gmail.com'
    gmail_pwd = 'Blackpearl@88'
    smtpserver = smtplib.SMTP("smtp.gmail.com", 587)
    smtpserver.ehlo()
    smtpserver.starttls()
    smtpserver.ehlo
    smtpserver.login(gmail_user,gmail_pwd)
    hostname = host
    comp=company
    surv=survey
    if ' ' in company:
        comp=company.replace(" ", "%20")
    if ' ' in survey:
        surv=survey.replace(" ", "%20")
    print(company, survey)
    for mail in to:

        subject ="Survey Request - {}".format(mail)
        smtpserver.login(gmail_user, gmail_pwd)
        header = "Reselience Mailer"

        text = "Dear member,\n" \
                   "\n" \
                   "\n" \
                   "you have been picked for taking the survey by {} and the link is a follows below.\n"\
                   "https://{}/backend/validation?company={}&survey={}&email={}&host={}".format(comp,hostname,comp, surv, mail, hostname)
        msg = 'Subject: {}\n\n{}'.format(subject,text)
        smtpserver.sendmail(gmail_user, mail, msg)

    smtpserver.close()





@app.route('/backend/getUsers', methods=['GET'])
def getUsers():

    company=request.args.get('company')
    dept=request.args.get('department')
    host,base,colection,dbuser,pwd=bl.mongoInit('userdetails')
    data=bl.getUsers(dept,company,host,base,colection,dbuser,pwd)
    return json.dumps(data)


@app.route('/backend/getUserId', methods=['GET'])
def getUserId():

    company=request.args.get('company')
    dept=request.args.get('department')
    email=request.args.get('email')
    survey=request.args.get('survey')
    host,base,colection,dbuser,pwd=bl.mongoInit('userdetails')

    data=bl.getID(dept,company,email,survey,host,base,colection,dbuser,pwd)
    id=''.join(str(data))

    return id

@app.route('/backend/questionsCount', methods=['GET'])
def questionsCount():
    company=request.args.get('company')
    survey=request.args.get('survey')
    count = bl.getquestioncount(company, survey, sectors, subsectors)
    return json.dumps(count, default=json_util.default)

@app.route('/backend/genReport', methods=['GET'])
def genReport():

    survey = request.args.get('survey')
    company = request.args.get('company')
    now = datetime.datetime.now()
    curdate=now.strftime("%Y-%m-%d")
    host,base,colection,dbuser,pwd=bl.mongoInit('users')
    #file='C:\\toHDD\\toViggu\\surveyapp_python\\dataframe_1555184982.xlsx'

    req = requests.get("http://localhost/backend/chartsAll?survey={}&company={}".format(survey,company))
    print(req.text)
    file=req.text

    r = requests.get("http://localhost/backend/getAllRadar?dataframe={}".format(file))
    jsonres=r.json()
    subscores=[]
    subscores.append(jsonres['Physical'][0])
    subscores.append(jsonres['Organizational'][0])
    subscores.append(jsonres['Technical'][0])
    subscores.append(jsonres['Physical'][1])
    subscores.append(jsonres['Organizational'][1])
    subscores.append(jsonres['Technical'][1])
    subscores.append(jsonres['Physical'][2])
    subscores.append(jsonres['Organizational'][2])
    subscores.append(jsonres['Technical'][2])
    subscores.append(jsonres['Physical'][3])
    subscores.append(jsonres['Organizational'][3])
    subscores.append(jsonres['Technical'][3])

    R0=cl.genFullRadar(sectors,subsectors,subscores)

    r = requests.get("http://localhost/backend/getAllRadarBySector?dataframe={}".format(file))
    jsonres=r.json()

    R1_labels=jsonres[0]['labels']
    R1_values=jsonres[0]['score']
    R2_labels=jsonres[1]['labels']
    R2_values=jsonres[1]['score']
    R3_labels=jsonres[2]['labels']
    R3_values=jsonres[2]['score']
    R4_labels=jsonres[3]['labels']
    R4_values=jsonres[3]['score']

    R1=cl.genRadar(R1_labels,R1_values,'#ff6666','Robustness')
    R2=cl.genRadar(R2_labels,R2_values,'#ffcc99','Redundancy')
    R3=cl.genRadar(R3_labels,R3_values,'#99ff99','Resourcefulness')
    R4=cl.genRadar(R4_labels,R4_values,'#66b3ff','Rapidity')

    r = requests.get("http://localhost/backend/getAllRadarAllSectors?dataframe={}".format(file))
    jsonres=r.json()
    labels=jsonres['category']
    values=jsonres['scores']
    print(values)
    labels=pd.Series(labels)
    values=pd.Series(values)
    R5=cl.barh(labels,values)

    r = requests.get("http://localhost/backend/getAllRadar?dataframe={}".format(file))
    jsonres=r.json()

    phy=jsonres['Physical']
    org=jsonres['Organizational']
    tech=jsonres['Technical']
    subsecs=['Physical','Organizational','Technical']
    bar1=[phy[0],org[0],tech[0]]
    bar2=[phy[1],org[1],tech[1]]
    bar3=[phy[2],org[2],tech[2]]
    bar4=[phy[3],org[3],tech[3]]

    R1_bar=cl.barRadar('R1',bar1,subsecs)
    R2_bar=cl.barRadar('R2',bar2,subsecs)
    R3_bar=cl.barRadar('R3',bar3,subsecs)
    R4_bar=cl.barRadar('R4',bar4,subsecs)

    lists=[R0,R1,R2,R3,R4,R5,R1_bar,R2_bar,R3_bar,R4_bar]
    print(lists)

    doc = DocxTemplate("sampleWord.docx")
    s1 = doc.new_subdoc()
    s1.add_picture(R1, width=Inches(7))
    s2 = doc.new_subdoc()
    s2.add_picture(R2, width=Inches(7))
    s3 = doc.new_subdoc()
    s3.add_picture(R3, width=Inches(7))
    s4 = doc.new_subdoc()
    s4.add_picture(R4, width=Inches(7))
    s5 = doc.new_subdoc()
    s5.add_picture(R0, width=Inches(7))
    s6 = doc.new_subdoc()
    s6.add_picture(R5, width=Inches(7))

    s7 = doc.new_subdoc()
    s7.add_picture(R1_bar, width=Inches(5))
    s8 = doc.new_subdoc()
    s8.add_picture(R2_bar, width=Inches(5))
    s9 = doc.new_subdoc()
    s9.add_picture(R3_bar, width=Inches(5))
    s10 = doc.new_subdoc()
    s10.add_picture(R4_bar, width=Inches(5))
    req = requests.get("http://localhost/backend/surveysCount?survey={}&company={}".format(survey,company))
    count=req.text
    
    context = { 'bar1' : s7,'bar2' : s8,'bar3' : s9,'bar4' : s10,'chart5' : s6, 'radar1' : s1,'radar2' : s2,'radar3' : s3,'radar4' : s4,'complete' : s5, 'company': company, 'survey': survey, 'date': curdate, 'number': count  }
    doc.render(context)
    filename="Survey_Metrics_{}.docx".format(bl.getTimeStamp())
    doc.save(filename)
    outputFile=bl.convertToPDF(filename)


    bl.sendReport(survey,company,outputFile)
    return "Sent Report"

@app.route('/backend/userDepartmentNames', methods=['GET', 'POST'])
def userDepartmentNames():
    company = request.args.get('company')
    survey = request.args.get('survey')
    #company = "company3"
    #survey = "survey3"
    departmentnames=list(bl.getuserdepartmentnames(company, survey))
    return json.dumps(departmentnames)

@app.route('/backend/userCompanyNames', methods=['GET', 'POST'])
def userCompanyNames():
    companies_list = list(bl.getusercompanynames())
    print(companies_list)
    return json.dumps(companies_list)

@app.route('/backend/userCompanySurveyNames', methods=['GET', 'POST'])
def userCompanySurveyNames():
    company = request.args.get('company')
    #company = "company3"
    companysurveynames=list(bl.getusercompanysurveynames(company))
    return json.dumps(companysurveynames)

@app.route('/backend/generateOTP', methods=['GET', 'POST'])
def generateOTP():
    email = request.args.get('email')
    company = request.args.get('company')
    print(email, company)
    mails = []
    #company = "company3"
    #email = "pranushajanapala@gmail.com"
    otp = bl.otp(email, company)
    mails.append(email)
    mail(company,mails,otp)
    return "done:"

def  mail(company,email,otp):
    app.config['MAIL_SERVER']='smtp.gmail.com'
    app.config['MAIL_PORT'] = 587
    app.config['MAIL_USERNAME'] = 'perugubharathkumar@gmail.com'
    app.config['MAIL_PASSWORD'] = 'Blackpearl@88'
    app.config['MAIL_USE_TLS'] = True
    app.config['MAIL_USE_SSL'] = False
    app.config['MAIL_DEFAULT_SENDER'] = 'pranushajanapala@gmail.com'
    mail=Mail(app)
    print(email)
    msg = Message("Password Reset Request", recipients=email)
    msg.body = "Dear member,\n" \
               "" \
               "" \
               "The otp to reset your password is - {}\n" \
               "Please provide the same at password reset page and proceed further".format(otp)
    mail.send(msg)

@app.route('/backend/passwordReset', methods=['GET', 'POST'])
def passwordReset():
    email = request.args.get('email')
    company = request.args.get('company')
    password =  request.args.get('password')
    forgotphrase =  request.args.get('forgotphrase')
    #company = "company3"
    #email = "pranushajanapala@gmail.com"
    #password = "password"
    #forgotphrase = "3944"
    return bl.resetpassword(email,company,password,forgotphrase)

@app.route('/backend/getAllSurveybyUser', methods=['GET', 'POST'])
def getAllSurveybyUser():
    email = request.args.get('email')
    company = request.args.get('company')
    return bl.getAllSurvey(email, company)

@app.route('/backend/encrypt', methods=['GET','POST'])
def encrypt():
    plain_text = request.args.get('str')
    en_de = request.args.get('enc')
    return bl.en_de_crypt(plain_text, en_de)  

@app.route('/backend/surveysCount', methods=['GET', 'POST'])
def surveysCount():
    survey = request.args.get('survey')
    company = request.args.get('company')
    return bl.getnoofsurveys(survey, company)


if __name__ == '__main__':
    app.run(host='0.0.0.0')